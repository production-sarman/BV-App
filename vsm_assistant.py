
import os
import json
import re
import shutil
import zipfile
import mammoth
from datetime import datetime

from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse, HTMLResponse
from pydantic import BaseModel
from docx import Document
from fastapi.middleware.cors import CORSMiddleware
from fastapi import APIRouter
from typing import Optional


router = APIRouter()

# ==========================
# PATH CONFIG
# ==========================
VSM_ROOT        = os.path.dirname(os.path.abspath(__file__))
CENTRAL_JSON    = os.path.join(VSM_ROOT, "central_index.json")
WORKING_COPIES  = os.path.join(VSM_ROOT, "Working_Copies")
VSM_DATA_FOLDER = os.path.join(VSM_ROOT, "vsm_data")
MASTER_DOCS     = os.path.join(VSM_ROOT, "Master_Documents")

os.makedirs(WORKING_COPIES, exist_ok=True)
os.makedirs(VSM_DATA_FOLDER, exist_ok=True)

# ==========================
# SUPPORTED FILE TYPES
# ==========================
SUPPORTED_FILE_TYPES = {
    ".docx":   ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", True),
    ".txt":    ("text/plain", True),
    ".pdf":    ("application/pdf", True),
    ".xls":    ("application/vnd.ms-excel", False),
    ".xlsx":   ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", False),
    ".sldprt": ("application/octet-stream", False),
    ".sldasm": ("application/octet-stream", False),
    ".step":   ("application/octet-stream", False),
    ".stp":    ("application/octet-stream", False),
}

FILE_PRIORITY = [".docx", ".pdf", ".txt", ".xlsx", ".xls", ".sldprt", ".sldasm", ".step", ".stp"]

# ==========================
# LOAD CENTRAL INDEX
# ==========================
if not os.path.exists(CENTRAL_JSON):
    raise FileNotFoundError("central_index.json missing")

with open(CENTRAL_JSON, "r", encoding="utf-8") as f:
    central_index = json.load(f)

# ==========================
# FILE HELPERS
# ==========================
def get_file_ext(path: str) -> str:
    return os.path.splitext(path)[1].lower()

def get_mime_type(path: str) -> str:
    mime, _ = SUPPORTED_FILE_TYPES.get(get_file_ext(path), ("application/octet-stream", False))
    return mime

def is_previewable(path: str) -> bool:
    _, previewable = SUPPORTED_FILE_TYPES.get(get_file_ext(path), ("application/octet-stream", False))
    return previewable

def is_supported(path: str) -> bool:
    return get_file_ext(path) in SUPPORTED_FILE_TYPES

def pick_best_file(documents: list) -> Optional[str]:
    existing = [p for p in documents if os.path.exists(p)]
    if not existing:
        return None
    for preferred_ext in FILE_PRIORITY:
        for path in existing:
            if get_file_ext(path) == preferred_ext:
                return path
    return existing[0]

def read_stage_explanation(stage: str) -> str:
    path = os.path.join(VSM_DATA_FOLDER, f"{stage}_explanation.docx")
    if not os.path.exists(path):
        return f"No explanation available for stage {stage}"
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def create_versioned_copy(source_path: str) -> str:
    date_folder = os.path.join(WORKING_COPIES, datetime.now().strftime("%Y-%m-%d"))
    os.makedirs(date_folder, exist_ok=True)
    base = os.path.basename(source_path)
    name, ext = os.path.splitext(base)
    version = 1
    while True:
        dest = os.path.join(date_folder, f"{name}_v{version}{ext}")
        if not os.path.exists(dest):
            break
        version += 1
    shutil.copy2(source_path, dest)
    return dest

def normalise_documents_field(stage_data: dict) -> list:
    if "documents" in stage_data:
        value = stage_data["documents"]
        return value if isinstance(value, list) else [value]
    if "document" in stage_data:
        value = stage_data["document"]
        return [value] if value else []
    return []

# ==========================
# ZIP HELPER
# Zips ALL files in subfolder regardless of extension
# ==========================
# ==========================
# ZIP HELPER
# Zips ALL files recursively — entire folder tree
# ==========================
def create_zip_of_folder(subfolder_path: str, subfolder_name: str) -> str:
    date_folder = os.path.join(WORKING_COPIES, datetime.now().strftime("%Y-%m-%d"))
    os.makedirs(date_folder, exist_ok=True)

    safe_name = re.sub(r"[^\w\-.]", "_", subfolder_name)
    version   = 1
    while True:
        zip_path = os.path.join(date_folder, f"{safe_name}_v{version}.zip")
        if not os.path.exists(zip_path):
            break
        version += 1

    files_added = 0
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(subfolder_path):
            for fname in files:
                fpath   = os.path.join(root, fname)
                arcname = os.path.relpath(fpath, subfolder_path)
                zf.write(fpath, arcname=arcname)
                files_added += 1

    print(f"  [ZIP] Created with {files_added} files: {os.path.basename(zip_path)}")
    return zip_path

# ==========================
# UNIT NORMALISATION
# ==========================
def normalize_units(text: str) -> str:
    def ka_repl(match):
        value = float(match.group(1))
        return str(int(value * 1000)) + "A"
    text = re.sub(r"(\d+(?:\.\d+)?)\s*[Kk][Aa]", ka_repl, text)
    text = re.sub(r"(\d+(?:\.\d+)?)-(\d+(?:\.\d+)?[A-Za-z]*)", r"\1~\2", text)
    text = re.sub(r"[\s\-_]+", "_", text.strip())
    text = text.replace("~", "-")
    return text.upper()

# ==========================
# FOLDER STRUCTURE SCANNER
# ==========================
def _norm(text: str) -> str:
    return re.sub(r"[\s\-_]+", "_", text.strip().upper())

def build_folder_lookup() -> dict:
    lookup = {}
    if not os.path.exists(MASTER_DOCS):
        return lookup

    for product_type in os.listdir(MASTER_DOCS):
        product_path = os.path.join(MASTER_DOCS, product_type)
        if not os.path.isdir(product_path):
            continue

        for variant in os.listdir(product_path):
            variant_path = os.path.join(product_path, variant)
            if not os.path.isdir(variant_path):
                continue

            subfolders = [
                d for d in os.listdir(variant_path)
                if os.path.isdir(os.path.join(variant_path, d))
            ]

            index_key    = normalize_units(variant.upper())
            norm_variant = _norm(index_key)

            alias_keys = set()
            alias_keys.add(norm_variant)
            alias_keys.add(_norm(variant))
            alias_keys.add(_norm(product_type))
           
            for sf in subfolders:
                alias_keys.add(f"{norm_variant}_{_norm(sf)}")
                alias_keys.add(f"{_norm(variant)}_{_norm(sf)}")
                alias_keys.add(f"{_norm(product_type)}_{_norm(sf)}")

            lookup[norm_variant] = {
                "variant_key":  index_key,
                "product_type": product_type,
                "subfolders":   [s.upper() for s in subfolders],
                "alias_keys":   list(alias_keys),
            }

    return lookup

FOLDER_LOOKUP: dict = build_folder_lookup()

# ==========================
# SMART QUERY PARSER
# ==========================
def extract_stage(text: str) -> str | None:
    match = re.search(r"\b(\d+\.\d+)\b", text)
    return match.group(1) if match else None

def clean_query(query: str) -> str:
    q = normalize_units(query)
    q = re.sub(r"\b\d+\.\d+\b", "", q)
    q = re.sub(
        r"\b(stage|for|of|get|give|me|the|show|fetch|find|please|document|file)\b",
        "", q, flags=re.I
    )
    return q.strip()

def tokenise(text: str) -> list:
    return [t.upper() for t in re.split(r"[\s\-_]+", text) if len(t) >= 2]

def score_match(query_tokens: list, candidate_norm_key: str) -> int:
    return sum(1 for t in query_tokens if t in candidate_norm_key)

def fuzzy_match_variant(query: str) -> str | None:
    cleaned      = clean_query(query)
    query_norm   = _norm(cleaned)
    query_tokens = tokenise(cleaned)

    if not query_tokens:
        return None

    # Pass 1: exact alias match
    for norm_variant, info in FOLDER_LOOKUP.items():
        for alias in info["alias_keys"]:
            if query_norm == alias:
                return info["variant_key"]

    # Pass 2: token-score against all aliases
    best_variant = None
    best_score   = 0

    for norm_variant, info in FOLDER_LOOKUP.items():
        all_keys_to_score = info["alias_keys"] + [norm_variant]
        for alias in all_keys_to_score:
            s = score_match(query_tokens, alias)
            if s > best_score:
                best_score   = s
                best_variant = info["variant_key"]

    threshold = max(1, len(query_tokens) // 2)
    if best_variant and best_score >= threshold:
        return best_variant

    # Pass 3: fallback against central_index keys directly
    all_index_variants = []
    for product_data in central_index.values():
        all_index_variants.extend(product_data.get("variants", {}).keys())

    best_variant = None
    best_score   = 0

    for vk in all_index_variants:
        s = score_match(query_tokens, _norm(vk))
        if s > best_score:
            best_score   = s
            best_variant = vk

    if best_variant and best_score >= threshold:
        return best_variant

    return None

def parse_query(user_input: str):
    stage       = extract_stage(user_input)
    variant_key = fuzzy_match_variant(user_input)
    return stage, variant_key

# ==========================
# REQUEST MODEL
# ==========================
class QueryRequest(BaseModel):
    query: str

# ==========================
# QUERY API
# ==========================
@router.post("/query")
def process_query(request: QueryRequest):

    user_input          = request.query.strip()
    stage, variant_code = parse_query(user_input)


#should i place ur code overhere

    if not stage:
        raise HTTPException(
            status_code=400,
            detail=(
                "Could not find a stage number in your query. "
                "Try: 'stage 1.3 TR_3PH_CS_01_6V_2000A' or '1.3 6V 2000A'"
            )
        )

    if not variant_code:
        all_variants = [
            vk
            for pd in central_index.values()
            for vk in pd.get("variants", {}).keys()
        ]
        raise HTTPException(
            status_code=404,
            detail=(
                f"Could not match any product from your query '{user_input}'. "
                f"Known products: {all_variants}"
            )
        )

    for product_type, product_data in central_index.items():
        variants = product_data.get("variants", {})

        if variant_code not in variants:
            continue

        variant_data     = variants[variant_code]
        explanation      = read_stage_explanation(stage)
        available_stages = list(variant_data.keys())

        if stage not in variant_data:
            return {
                "stage":            stage,
                "product_code":     variant_code,
                "explanation":      f"Stage {stage} not available. Available: {available_stages}",
                "download_url":     "",
                "preview_url":      "",
                "available_stages": available_stages,
                "file_type":        None,
                "all_file_types":   [],
            }

        stage_data = variant_data[stage]
        entry_type = stage_data.get("type", "file")

        # ==========================
        # SUBFOLDER ENTRY
        # Always ZIP entire folder — all file types included
        # ==========================
        if entry_type == "subfolder":
            subfolder_path = stage_data.get("subfolder_path", "")
            subfolder_name = stage_data.get("subfolder_name", "subfolder")

            if not subfolder_path or not os.path.exists(subfolder_path):
                return {
                    "stage":            stage,
                    "product_code":     variant_code,
                    "explanation":      explanation,
                    "download_url":     "",
                    "preview_url":      "",
                    "available_stages": available_stages,
                    "file_type":        "folder",
                    "all_file_types":   [],
                    "error":            f"Subfolder not found on disk: {subfolder_path}",
                }

            # Count ALL files recursively — not just root level
            actual_files = []
            actual_exts  = set()
            for root, dirs, files in os.walk(subfolder_path):
                for fname in files:
                    actual_files.append(fname)
                    actual_exts.add(get_file_ext(fname))

            if not actual_files:
                return {
                    "stage":            stage,
                    "product_code":     variant_code,
                    "explanation":      explanation,
                    "download_url":     "",
                    "preview_url":      "",
                    "available_stages": available_stages,
                    "file_type":        "folder",
                    "all_file_types":   [],
                    "error":            "Subfolder exists but contains no files.",
                }

            # Create ZIP of entire folder tree recursively
            zip_path = create_zip_of_folder(subfolder_path, subfolder_name)
            relative = os.path.relpath(zip_path, WORKING_COPIES).replace("\\", "/")

            return {
                "stage":              stage,
                "product_code":       variant_code,
                "matched_from_query": user_input,
                "explanation":        explanation,
                "download_url":       f"/download_doc/?relative_path={relative}",
                "preview_url":        "",
                "available_stages":   available_stages,
                "file_type":          ".zip",
                "all_file_types":     list(actual_exts),
                "entry_type":         "subfolder_zip",
                "subfolder_name":     subfolder_name,
                "total_files_in_zip": len(actual_files),
            }
        # ==========================
        # NORMAL FILE ENTRY
        # Existing behaviour — unchanged
        # ==========================
        all_paths = normalise_documents_field(stage_data)

        if not all_paths:   
            return {
                "stage":            stage,
                "product_code":     variant_code,
                "explanation":      explanation,
                "download_url":     "",
                "preview_url":      "",
                "available_stages": available_stages,
                "file_type":        None,
                "all_file_types":   [],
            }

    best_path = pick_best_file(all_paths)

    # ==========================
    # ✅ FIRST: check best_path
    # ==========================
    if best_path is None:
        return {
            "stage": stage,
            "product_code": variant_code,
            "explanation": explanation,
            "download_url": "",
            "preview_url": "",
            "available_stages": available_stages,
            "file_type": None,
            "all_file_types": [],
            "error": "No files found on disk for this stage.",
        }

    # ==========================
    # 🔥 DYNAMIC FOLDER ZIP LOGIC
    # ==========================
    if stage_data.get("type") == "file" and stage_data.get("source_folder"):
        folder_path = stage_data.get("source_folder")

        if not os.path.exists(folder_path):
            return {
                "stage": stage,
                "product_code": variant_code,
                "error": f"Folder not found: {folder_path}"
            }

        zip_path = create_zip_of_folder(folder_path, f"{variant_code}_{stage}")
        relative = os.path.relpath(zip_path, WORKING_COPIES).replace("\\", "/")

        return {
            "stage": stage,
            "product_code": variant_code,
            "matched_from_query": user_input,
            "explanation": explanation,
            "download_url": f"/download_doc/?relative_path={relative}",
            "preview_url": "",
            "available_stages": available_stages,
            "file_type": ".zip",
            "entry_type": "dynamic_folder_zip",
        }

    # ==========================
    # ✅ EXISTING LOGIC (UNCHANGED)
    # ==========================
    if not is_supported(best_path):
        return {
            "stage": stage,
            "product_code": variant_code,
            "explanation": explanation,
            "download_url": "",
            "preview_url": "",
            "available_stages": available_stages,
            "file_type": get_file_ext(best_path),
            "all_file_types": [get_file_ext(p) for p in all_paths if os.path.exists(p)],
            "error": f"Unsupported file type: {get_file_ext(best_path)}",
        }

    working_path  = create_versioned_copy(best_path)
    relative_path = os.path.relpath(working_path, WORKING_COPIES).replace("\\", "/")
    ext           = get_file_ext(best_path)

    return {
        "stage": stage,
        "product_code": variant_code,
        "matched_from_query": user_input,
        "explanation": explanation,
        "download_url": f"/download_doc/?relative_path={relative_path}",
        "preview_url": (
            f"/view_doc/?relative_path={relative_path}"
            if is_previewable(best_path) else ""
        ),
        "available_stages": available_stages,
        "file_type": ext,
        "all_file_types": [get_file_ext(p) for p in all_paths if os.path.exists(p)],
        "entry_type": "file",
    }


# ==========================
# RELOAD INDEX
# ==========================
@router.post("/reload")
def reload_index():
    global central_index, FOLDER_LOOKUP

    if not os.path.exists(CENTRAL_JSON):
        raise HTTPException(status_code=500, detail="central_index.json missing")

    with open(CENTRAL_JSON, "r", encoding="utf-8") as f:
        central_index = json.load(f)

    FOLDER_LOOKUP = build_folder_lookup()

    return { 
        "message":            "Index reloaded successfully",
        "products":           len(central_index),
        "variants_in_products": len(FOLDER_LOOKUP)

    }


# ==========================
# DOWNLOAD DOC (Universal — works for ZIP too)
# ==========================
@router.get("/download_doc/")
def download_doc(relative_path: str):
    file_path = os.path.abspath(os.path.join(WORKING_COPIES, relative_path))
    base      = os.path.abspath(WORKING_COPIES)

    if not file_path.startswith(base):
        raise HTTPException(status_code=403, detail="Access denied")
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    

    ext = get_file_ext(file_path)
    if ext == ".zip":
        mime = "application/zip"
    else:
        mime = get_mime_type(file_path)

    return FileResponse(
        path=file_path,
        filename=os.path.basename(file_path),
        media_type=mime,
    )


# ==========================
# PREVIEW DOC
# ==========================
@router.get("/view_doc/", response_class=HTMLResponse)
def view_doc(relative_path: str):
    file_path = os.path.abspath(os.path.join(WORKING_COPIES, relative_path))
    base      = os.path.abspath(WORKING_COPIES)

    if not file_path.startswith(base):
        raise HTTPException(status_code=403, detail="Access denied")
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    ext = get_file_ext(file_path)

    if ext == ".docx":
        with open(file_path, "rb") as f:
            result = mammoth.convert_to_html(f)
        return f"<html><body style='font-family:Arial;padding:40px;'>{result.value}</body></html>"

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        content_escaped = (
            content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        )
        return f"<html><body style='font-family:monospace;padding:40px;white-space:pre-wrap;'>{content_escaped}</body></html>"

    if ext == ".pdf":
        return FileResponse(
            path=file_path,
            filename=os.path.basename(file_path),
            media_type="application/pdf",
            headers={"Content-Disposition": "inline"},
        )

    raise HTTPException(
        status_code=415,
        detail=f"Preview not available for '{ext}'. Use the download URL instead.",
    )



# ==========================
# IMPROVED REPLACE VALUES (Preserves Alignment & Images)
# ==========================
@router.post("/replace_values")
def replace_values(relative_path: str, replacements: dict):
    """
    Replaces text while keeping images and formatting intact by 
    iterating through runs instead of overwriting the whole paragraph.
    """
    file_path = os.path.abspath(os.path.join(WORKING_COPIES, relative_path))
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    if get_file_ext(file_path) != ".docx":
        raise HTTPException(status_code=415, detail="Only .docx supported.")

    doc = Document(file_path)
    
    for p in doc.paragraphs:
        for run in p.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, str(new))
                    
    # Handle tables as well (common cause of alignment issues)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for old, new in replacements.items():
                            if old in run.text:
                                run.text = run.text.replace(old, str(new))

    doc.save(file_path)
    return {"message": "Updated successfully while preserving layout"}

# ==========================
# NEW: BROWSER-BASED EDIT & SAVE
# ==========================
class EditRequest(BaseModel):
    relative_path: str
    html_content: str

@router.post("/save_edit")
def save_edit(request: EditRequest):
    """
    Note: Converting HTML back to DOCX is risky for alignment.
    For high-precision, we recommend using the replace_values endpoint.
    This is a fallback for general text edits.
    """
    file_path = os.path.abspath(os.path.join(WORKING_COPIES, request.relative_path))
    
    # Logic to convert HTML back to DOCX would go here.
    # To keep 100% alignment, it's safer to use the 'replacements' 
    # method above via a frontend form.
    return {"message": "Content received. Use replace_values for image safety."}

# ==========================
# UPDATED VIEW DOC (With simple Edit UI)
# ==========================
@router.get("/view_doc/", response_class=HTMLResponse)
def view_doc(relative_path: str):
    file_path = os.path.abspath(os.path.join(WORKING_COPIES, relative_path))
    base      = os.path.abspath(WORKING_COPIES)

    # ✅ SECURITY FIX (was missing in your second version)
    if not file_path.startswith(base):
        raise HTTPException(status_code=403, detail="Access denied")

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    ext = get_file_ext(file_path)

    # ==========================
    # DOCX PREVIEW (SAFE)
    # ==========================
    if ext == ".docx":
        with open(file_path, "rb") as f:
            result = mammoth.convert_to_html(f)
            html_body = result.value

        return f"""
        <html>
            <head>
                <style>
                    body {{ font-family: sans-serif; padding: 20px; line-height: 1.6; }}
                    .editor-container {{ border: 1px solid #ccc; padding: 20px; margin-bottom: 20px; }}
                    .toolbar {{ background: #f4f4f4; padding: 10px; border-radius: 5px; margin-bottom: 10px; }}
                    button {{ padding: 10px 20px; cursor: pointer; background: #007bff; color: white; border: none; border-radius: 3px; }}
                </style>
            </head>
            <body>
                <div class="toolbar">
                    <strong>Preview Mode</strong> (Images preserved in original file)
                </div>
                <div class="editor-container" contenteditable="false">
                    {html_body}
                </div>
                <p><small>To edit without impacting alignment, use the \"Replace Values\" form.</small></p>
            </body>
        </html>
        """

    # ==========================
    # TXT PREVIEW
    # ==========================
    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        content_escaped = (
            content.replace("&", "&amp;")
            content.replace("<", "&lt;")                                                                        
            content.replace(">", "&gt;")
        )

        return f"""
        <html>
            <body style='font-family:monospace;padding:40px;white-space:pre-wrap;'>
                {content_escaped}
            </body>
        </html>
        """

    # ==========================
    # PDF INLINE VIEW
    # ==========================
    if ext == ".pdf":
        return FileResponse(
            path=file_path,
            filename=os.path.basename(file_path),
            media_type="application/pdf",
            headers={"Content-Disposition": "inline"},
        )

    # ==========================
    # NOT SUPPORTED
    # ==========================
    raise HTTPException(
        status_code=415,
        detail=f"Preview not available for '{ext}'. Use download instead."
    )
 